import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def get_meses_ordenados():
    return ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio', 'Agosto', 'Septiembre', 'Octubre',
            'Noviembre', 'Diciembre']


def traducir_mes(mes):
    traduccion = {
        'January': 'Enero',
        'February': 'Febrero',
        'March': 'Marzo',
        'April': 'Abril',
        'May': 'Mayo',
        'June': 'Junio',
        'July': 'Julio',
        'August': 'Agosto',
        'September': 'Septiembre',
        'October': 'Octubre',
        'November': 'Noviembre',
        'December': 'Diciembre'
    }
    if pd.isna(mes):
        return 'Mes desconocido'
    return traduccion.get(mes, 'Mes desconocido')


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generar_docx_informe(
        ruta_docx,
        mes_seleccionado,
        info_camaras_selected,
        valores_infracciones_selected,
        resumen_general_selected,
        infracciones_por_camara_selected,
        canales_de_pago_selected,
        movimientos_juzgados_selected,
        actas_juzgados_altas_selected,
        actas_juzgados_bajas_selected,
        dataframe_lotes,
):
    doc = Document()

    # =====================================================
    # I. EQUIPAMIENTO
    # =====================================================
    doc.add_heading(
        "1. Detalle del Equipamiento Tecnológico Utilizado para el Servicio",
        level=1
    )

    doc.add_paragraph(
        "Se procede a especificar las características tecnológicas de los equipos empleados, incluyendo el número de serie, la descripción de la cámara, la marca y el modelo correspondiente.También se proporciona información detallada sobre el uso del equipamiento, ubicación, fechas de inicio, tipo de montaje, cantidad de equipos por punto de ubicación y el juzgado responsable. En cuanto al uso del equipamiento, se detallará su propósito específico para su correcta operación. La ubicación del equipamiento se especificará claramente, indicando su posición física exacta. Se proporcionarán mapas o diagramas detallados para facilitar su localización. Las fechas de inicio de la actividad se registrarán meticulosamente, junto con cualquier información relevante. El tipo de montaje necesario para el equipamiento se describirá en función de las especificaciones técnicas de las cámaras."
        ""
    )

    # 1.1
    doc.add_heading("1.1 Descripción del equipo", level=2)
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Número de serie"
    hdr[1].text = "Marca"
    hdr[2].text = "Modelo"

    for _, row in info_camaras_selected.iterrows():
        c = tabla.add_row().cells
        c[0].text = str(row.get("Número de serie", ""))
        c[1].text = str(row.get("Marca", ""))
        c[2].text = str(row.get("Modelo", ""))

    # 1.2
    doc.add_heading("1.2 Estado y características de los equipos", level=2)
    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Ubicación"
    hdr[1].text = "Tipo de montaje"
    hdr[2].text = "Inicio de actividad"
    hdr[3].text = "Número de cámaras"
    hdr[4].text = "Tipo Infracción"

    for _, row in info_camaras_selected.iterrows():
        c = tabla.add_row().cells
        c[0].text = str(row.get("Ubicación", ""))
        c[1].text = str(row.get("Tipo de montaje", ""))
        c[2].text = str(row.get("Inicio de actividad", ""))
        c[3].text = str(row.get("Número de cámaras", ""))
        c[4].text = str(row.get("Tipo Infraccion", ""))

    # =====================================================
    # II. CUADRO TARIFARIO
    # =====================================================
    doc.add_heading(
        f"II. Clasificación de las infracciones de Tránsito / Cuadro Tarifario {mes_seleccionado} 2025",
        level=1
    )

    doc.add_paragraph(
        "El informe que se presenta a continuación tiene como objetivo exponer una tabla que detalla "
        "la clasificación de las infracciones de tránsito y su correspondiente cuadro tarifario."
    )

    df = valores_infracciones_selected.drop(columns=["Año", "Mes"], errors="ignore")
    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = col

    for _, row in df.iterrows():
        c = tabla.add_row().cells
        for i, col in enumerate(df.columns):
            c[i].text = str(row[col])

    # =====================================================
    # III. PROCESAMIENTO
    # =====================================================
    doc.add_heading("III. Procesamiento de infracciones", level=1)

    doc.add_paragraph(
        "El equipo técnico de la Universidad Nacional de La Matanza ha desplegado un conjunto integral de procedimientos de adquisición de imágenes, diseñados para optimizar el funcionamiento del sistema de registro de infracciones de tránsito.Estos procedimientos abarcan una variedad de tecnologías, cámaras de diferente tipo de infracción instaladas en semáforos.En paralelo, se ha establecido un sólido sistema de registro de infracciones de tránsito que se integra perfectamente con los diversos mecanismos de adquisición de imágenes. Este sistema se caracteriza por su precisión y confiabilidad en la recopilación de datos, lo que permite una gestión eficiente de las infracciones viales.Para garantizar la integridad y precisión de este proceso, se ha implementado un proceso de auditoría que abarca todas las etapas del sistema. Esta auditoría comprende la verificación meticulosa de la calibración y funcionamiento de los dispositivos de captura de imágenes, así como una revisión detallada de los datos recopilados. Además, se lleva a cabo una validación rigurosa de la información registrada, junto con la detección y corrección de posibles discrepancias o irregularidades.Gracias a este enfoque integral de fiscalización, se han alcanzado resultados notables en términos de mejora continua y eficacia operativa. La combinación de diversos procedimientos de adquisición de imágenes, un sistema de registro de infracciones de tránsito bien estructurado y un proceso detallado de auditoría ha sido fundamental para lograr resultados destacados en la gestión del tráfico y la seguridad vial en los municipios y sus alrededores."
    )

    doc.add_heading("III.I Descripción de las etapas del procesamiento de lo producido", level=2)

    df = resumen_general_selected.drop(
        columns=["Año", "Mes", "Actas Fehacientes", "Valor UF"],
        errors="ignore"
    )

    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = col

    for _, row in df.iterrows():
        c = tabla.add_row().cells
        for i, col in enumerate(df.columns):
            c[i].text = str(row[col])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(f"SeabornPlots/resumen_general_{mes_seleccionado}.png", width=Inches(6))

    doc.add_paragraph("")
    # =====================================================
    # V. EFECTIVIDAD RECAUDATORIA
    # =====================================================
    doc.add_heading(
        f"V. Efectividad Recaudatoria {mes_seleccionado} 2025 Berisso",
        level=1
    )
    doc.add_paragraph(
        "En relación con la generación de ingresos, es importante destacar que los pagos realizados por la cancelación de las infracciones registradas por el sistema de captación de la Universidad Nacional de La Matanza están sujetos a una serie de procedimientos de control para garantizar su adecuada gestión, como así también la integración con los diferentes sistemas de gestión de infracciones Nacionales y Provinciales. Estos procedimientos se traducen en una recopilación detallada de los montos recaudados por el Municipio, los cuales se presentan a continuación en forma de cuadros y gráficos para una mejor comprensión y análisis.A continuación, se detallan los montos recaudados por el Municipio, según los diferentes procedimientos de control aplicados a los pagos efectuados por la cancelación de las infracciones registradas por el sistema de captación de la Universidad Nacional de La Matanza.Estos datos reflejan la eficacia de los procedimientos de control implementados en la gestión de los pagos por cancelación de infracciones, proporcionando una visión clara y precisa de los ingresos generados. La transparencia en la presentación de esta información es fundamental para una adecuada rendición de cuentas y una gestión financiera responsable por parte del Municipio y la Universidad Nacional de la Matanza"

    )
    doc.add_heading(
        "V.I Producción por punto de Infracción para cámaras de tipo Semáforo",
        level=2
    )

    df = infracciones_por_camara_selected.drop(
        columns=["Año", "Mes", "Valor uf"],
        errors="ignore"
    )

    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = col

    for _, row in df.iterrows():
        c = tabla.add_row().cells
        for i, col in enumerate(df.columns):
            c[i].text = str(row[col])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(f"SeabornPlots/actas_por_infraccion_{mes_seleccionado}.png", width=Inches(6))

    # =====================================================
    # VI. MEDIOS DE PAGO
    # =====================================================
    doc.add_heading(
        f"VI. Detalle de los medios de pagos utilizados - {mes_seleccionado} 2025",
        level=1
    )

    df = canales_de_pago_selected.drop(
        columns=[
            "Año", "Mes", "Sugit s/ ga",
            "Desglose gastos administ bocas de recaudación",
            "Desglose gastos administ sugit"
        ],
        errors="ignore"
    )

    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = col

    for _, row in df.iterrows():
        c = tabla.add_row().cells
        for i, col in enumerate(df.columns):
            c[i].text = str(row[col])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(f"SeabornPlots/piechart_pagos_{mes_seleccionado}.png", width=Inches(6))

    # =====================================================
    # VII. JUZGADOS
    # =====================================================
    doc.add_heading(
        f"VII. Detalles de actas del mes de {mes_seleccionado} 2025",
        level=1
    )

    doc.add_heading(
        f"VII.I Detalle de Actividad de los Juzgados",
        level=2
    )

    df = movimientos_juzgados_selected.drop(
        columns=["Año", "Mes"],
        errors="ignore"
    )

    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = col

    for _, row in df.iterrows():
        c = tabla.add_row().cells
        for i, col in enumerate(df.columns):
            c[i].text = str(row[col])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(f"SeabornPlots/actividad_juzgados_{mes_seleccionado}.png", width=Inches(6))
    # VII.II SUGIT altas
    doc.add_heading("VII.II Detalle del SUGIT", level=2)
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Año"
    hdr[1].text = "Mes"
    hdr[2].text = "Cantidad de actas"

    for _, row in actas_juzgados_altas_selected.iterrows():
        c = tabla.add_row().cells
        c[0].text = str(row["Año"])
        c[1].text = str(row["Mes"])
        c[2].text = str(row["Cantidad de actas"])

    # VII.III SUGIT bajas
    doc.add_heading("VII.III Detalle del SUGIT bajas", level=2)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Año"
    hdr[1].text = "Mes"

    for _, row in actas_juzgados_bajas_selected.iterrows():
        c = tabla.add_row().cells
        c[0].text = str(row["Año"])
        c[1].text = str(row["Mes"])

    # =====================================================
    # VIII. NOTIFICACIONES
    # =====================================================
    doc.add_heading("VIII. Detalle de Notificaciones", level=1)

    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Lote"
    hdr[1].text = "Fecha Impresión"
    hdr[2].text = "Fecha Vencimiento"
    hdr[3].text = "Cantidad de Actas"

    for _, row in dataframe_lotes.iterrows():
        c = tabla.add_row().cells
        c[0].text = str(row["Lote"])
        c[1].text = str(row["Fecha Impresión"])
        c[2].text = str(row["Fecha Vencimiento"])
        c[3].text = str(row["Cantidad de Actas"])

    doc.add_paragraph(
        f"VIII.I Notificaciones realizadas en el mes de {mes_seleccionado}: "
        f"{int(dataframe_lotes['Cantidad de Actas'].sum())}"
    )

    # =====================================================
    # GUARDAR
    # =====================================================
    doc.save(ruta_docx)
